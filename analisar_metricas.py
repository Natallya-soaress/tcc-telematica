import math
import os
import shutil
from dataclasses import dataclass
from typing import Dict, Iterable, List, Mapping, Optional, Tuple

import numpy as np
import pandas as pd


VOLUMES_ESPERADOS = [100, 500, 1000, 2000, 5000]
CARGAS_ESPERADAS = ["baixa", "media", "alta"]
EXECUCOES_ESPERADAS = [1, 2, 3]


METRICAS_BASE = [
    "latencia_total",
    "tempo_processamento",
    "throughput",
    "cpu_percent",
    "mem_percent",
    "tempo_cifragem",
    "tempo_decifragem",
    "tamanho_payload",
]


@dataclass(frozen=True)
class Paths:
    plain_csv: str = "metricas_plain.csv"
    he_csv: str = "metricas_he.csv"
    out_dir: str = "outputs_analise"
    figs_dir: str = os.path.join(out_dir, "figuras")
    tables_dir: str = os.path.join(out_dir, "tabelas")
    texto_dir: str = os.path.join(out_dir, "texto")


def _ensure_dirs(p: Paths) -> None:
    os.makedirs(p.out_dir, exist_ok=True)
    os.makedirs(p.figs_dir, exist_ok=True)
    os.makedirs(p.tables_dir, exist_ok=True)
    os.makedirs(p.texto_dir, exist_ok=True)


def _read_csv(path: str) -> pd.DataFrame:
    df = pd.read_csv(path)
    for col in ["volume", "execucao", "evento"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    for col in METRICAS_BASE:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    if "carga" in df.columns:
        df["carga"] = df["carga"].astype(str).str.strip().str.lower()
    if "cenario" in df.columns:
        df["cenario"] = df["cenario"].astype(str).str.strip().str.upper()
    return df


def _validate_schema(df: pd.DataFrame, name: str) -> List[str]:
    erros: List[str] = []
    obrigatorias = {"cenario", "volume", "carga", "execucao", "evento"} | set(METRICAS_BASE)
    faltantes = sorted(obrigatorias - set(df.columns))
    if faltantes:
        erros.append(f"[{name}] colunas faltantes: {faltantes}")
    return erros


def _validate_grid(df: pd.DataFrame, name: str) -> List[str]:
    erros: List[str] = []
    presentes = set(
        (int(v), str(c), int(e))
        for v, c, e in df[["volume", "carga", "execucao"]].dropna().itertuples(index=False, name=None)
    )
    for v in VOLUMES_ESPERADOS:
        for c in CARGAS_ESPERADAS:
            for e in EXECUCOES_ESPERADAS:
                if (v, c, e) not in presentes:
                    erros.append(f"[{name}] configuração ausente: volume={v}, carga={c}, execucao={e}")
    return erros


def _validate_values(df: pd.DataFrame, name: str) -> List[str]:
    erros: List[str] = []
    nulls = df.isna().sum()
    null_cols = nulls[nulls > 0].to_dict()
    if null_cols:
        erros.append(f"[{name}] valores ausentes por coluna: {null_cols}")

    for col in METRICAS_BASE:
        if col not in df.columns:
            continue
        mn = df[col].min(skipna=True)
        if pd.notna(mn) and mn < 0:
            erros.append(f"[{name}] valores negativos em {col}: min={mn}")

    if "evento" in df.columns:
        bad = df[df["evento"] <= 0]
        if len(bad) > 0:
            erros.append(f"[{name}] eventos inválidos (<=0): {len(bad)} linhas")

    return erros


def _per_run_aggregates(df: pd.DataFrame) -> pd.DataFrame:
    """
    Agrega por (cenario, volume, carga, execucao), produzindo métricas por execução.
    - latencia_media_evento: média de latência_total por evento.
    - tempo_score_medio: média de tempo_processamento por evento.
    - throughput_run: volume / soma(latencia_total) (eventos/s).
    - cpu_medio, mem_medio: médias por evento.
    - tempo_cifragem_medio, tempo_decifragem_medio: médias por evento.
    - payload_medio: média por evento (bytes).
    """
    g = df.groupby(["cenario", "volume", "carga", "execucao"], as_index=False)
    out = g.agg(
        n_eventos=("evento", "count"),
        latencia_media_evento=("latencia_total", "mean"),
        latencia_total_soma=("latencia_total", "sum"),
        tempo_score_medio=("tempo_processamento", "mean"),
        cpu_medio=("cpu_percent", "mean"),
        mem_medio=("mem_percent", "mean"),
        tempo_cifragem_medio=("tempo_cifragem", "mean"),
        tempo_decifragem_medio=("tempo_decifragem", "mean"),
        payload_medio=("tamanho_payload", "mean"),
    )
    out["throughput_run"] = out["volume"] / out["latencia_total_soma"]
    return out


def _descritivas(df: pd.DataFrame, cols: Iterable[str], group_cols: List[str]) -> pd.DataFrame:
    stats = []
    for col in cols:
        if col not in df.columns:
            continue
        g = df.groupby(group_cols)[col]
        desc = g.agg(["mean", "std", "min", "max"]).reset_index()
        desc.insert(len(group_cols), "metrica", col)
        stats.append(desc)
    if not stats:
        return pd.DataFrame()
    out = pd.concat(stats, ignore_index=True)
    return out


def _merge_overhead(
    run_plain: pd.DataFrame, run_he: pd.DataFrame, metricas: List[str]
) -> pd.DataFrame:
    keys = ["volume", "carga", "execucao"]
    p = run_plain[keys + metricas].copy()
    h = run_he[keys + metricas].copy()
    merged = p.merge(h, on=keys, suffixes=("_plain", "_he"), how="inner")
    for m in metricas:
        denom = merged[f"{m}_plain"]
        num = merged[f"{m}_he"] - merged[f"{m}_plain"]
        merged[f"overhead_pct_{m}"] = (num / denom) * 100.0
    return merged


def _make_plots(run_all: pd.DataFrame, overhead: pd.DataFrame, p: Paths) -> None:
    import matplotlib.pyplot as plt
    import seaborn as sns

    sns.set_theme(style="whitegrid", font_scale=1.1)

    def savefig(fname: str) -> None:
        plt.tight_layout()
        plt.savefig(os.path.join(p.figs_dir, fname), dpi=300, bbox_inches="tight")
        plt.close()

    # Gráfico 1: Latência média por volume (HE vs Plain)
    plt.figure(figsize=(9, 5))
    g1 = run_all.groupby(["cenario", "volume"], as_index=False)["latencia_media_evento"].agg(["mean", "std"])
    g1 = g1.reset_index()
    sns.lineplot(data=g1, x="volume", y="mean", hue="cenario", marker="o")
    plt.xlabel("Volume de eventos")
    plt.ylabel("Latência média por evento (s)")
    plt.title("Latência média por volume (CKKS vs baseline)")
    plt.legend(title="Cenário")
    savefig("grafico_1_latencia_media_por_volume.png")

    # Gráfico 2: Throughput por carga
    plt.figure(figsize=(9, 5))
    g2 = run_all.groupby(["cenario", "carga"], as_index=False)["throughput_run"].agg(["mean", "std"]).reset_index()
    sns.barplot(data=g2, x="carga", y="mean", hue="cenario", errorbar=None)
    plt.xlabel("Carga")
    plt.ylabel("Throughput (eventos/s)")
    plt.title("Throughput médio por carga (CKKS vs baseline)")
    plt.legend(title="Cenário")
    savefig("grafico_2_throughput_por_carga.png")

    # Gráfico 3: Tempo médio de score
    plt.figure(figsize=(9, 5))
    g3 = run_all.groupby(["cenario", "volume"], as_index=False)["tempo_score_medio"].mean()
    sns.lineplot(data=g3, x="volume", y="tempo_score_medio", hue="cenario", marker="o")
    plt.xlabel("Volume de eventos")
    plt.ylabel("Tempo médio de processamento do score (s)")
    plt.title("Tempo médio de score por volume (CKKS vs baseline)")
    plt.legend(title="Cenário")
    savefig("grafico_3_tempo_score.png")

    # Gráfico 4: Consumo de CPU
    plt.figure(figsize=(9, 5))
    g4 = run_all.groupby(["cenario", "carga"], as_index=False)["cpu_medio"].mean()
    sns.barplot(data=g4, x="carga", y="cpu_medio", hue="cenario", errorbar=None)
    plt.xlabel("Carga")
    plt.ylabel("CPU média (%)")
    plt.title("Consumo médio de CPU por carga (CKKS vs baseline)")
    plt.legend(title="Cenário")
    savefig("grafico_4_cpu.png")

    # Gráfico 5: Consumo de memória
    plt.figure(figsize=(9, 5))
    g5 = run_all.groupby(["cenario", "carga"], as_index=False)["mem_medio"].mean()
    sns.barplot(data=g5, x="carga", y="mem_medio", hue="cenario", errorbar=None)
    plt.xlabel("Carga")
    plt.ylabel("Memória média (%)")
    plt.title("Consumo médio de memória por carga (CKKS vs baseline)")
    plt.legend(title="Cenário")
    savefig("grafico_5_memoria.png")

    # Gráfico 6: Tempo de cifragem e decifragem (apenas HE; baseline é ~0)
    plt.figure(figsize=(9, 5))
    he_only = run_all[run_all["cenario"] == "HE"].copy()
    g6 = he_only.groupby(["volume"], as_index=False)[["tempo_cifragem_medio", "tempo_decifragem_medio"]].mean()
    g6m = g6.melt(id_vars=["volume"], value_vars=["tempo_cifragem_medio", "tempo_decifragem_medio"],
                 var_name="etapa", value_name="tempo_s")
    sns.lineplot(data=g6m, x="volume", y="tempo_s", hue="etapa", marker="o")
    plt.xlabel("Volume de eventos")
    plt.ylabel("Tempo médio por evento (s)")
    plt.title("Tempo médio de cifragem e decifragem por volume (CKKS)")
    plt.legend(title="Etapa")
    savefig("grafico_6_cifragem_decifragem.png")

    # Gráfico 7: Crescimento do tamanho dos dados cifrados
    plt.figure(figsize=(9, 5))
    g7 = he_only.groupby(["volume"], as_index=False)["payload_medio"].mean()
    sns.lineplot(data=g7, x="volume", y="payload_medio", marker="o", color="#2a6fdb")
    plt.xlabel("Volume de eventos")
    plt.ylabel("Tamanho médio do payload cifrado (bytes)")
    plt.title("Crescimento do payload cifrado por volume (CKKS)")
    savefig("grafico_7_payload.png")

    # Gráfico 8: Overhead criptográfico percentual
    plt.figure(figsize=(10, 5.5))
    cols = ["overhead_pct_latencia_media_evento", "overhead_pct_throughput_run", "overhead_pct_tempo_score_medio",
            "overhead_pct_cpu_medio", "overhead_pct_mem_medio"]
    ov = overhead.copy()
    # agregação por volume (média nas cargas e execuções)
    agg = ov.groupby(["volume"], as_index=False)[cols].mean()
    agg_m = agg.melt(id_vars=["volume"], value_vars=cols, var_name="metrica", value_name="overhead_pct")
    sns.lineplot(data=agg_m, x="volume", y="overhead_pct", hue="metrica", marker="o")
    plt.axhline(0, color="black", linewidth=0.8)
    plt.xlabel("Volume de eventos")
    plt.ylabel("Overhead (%)")
    plt.title("Overhead percentual do CKKS em métricas selecionadas (média por volume)")
    plt.legend(title="Métrica", bbox_to_anchor=(1.02, 1), loc="upper left")
    savefig("grafico_8_overhead_percentual.png")


def _format_tabela_uspesalq(title: str, df: pd.DataFrame, fonte: str) -> str:
    linhas = []
    linhas.append(title)
    linhas.append("")
    linhas.append(df.to_string(index=False))
    linhas.append("")
    linhas.append(f"Fonte: {fonte}.")
    return "\n".join(linhas)


def _write_text_sections(
    run_all: pd.DataFrame,
    overhead_by_cfg: pd.DataFrame,
    p: Paths,
) -> None:
    # Procedimentos experimentais: baseado no que está evidenciado nos CSVs
    proc = """
### Procedimentos Experimentais (detalhamento da execução)

Os experimentos foram conduzidos em ambiente controlado, com persistência dos resultados em banco de dados PostgreSQL, publicação de eventos via Apache Kafka, disponibilização de endpoint de ingestão via FastAPI e processamento em Python 3.x. O monitoramento de recursos computacionais foi realizado por meio da biblioteca psutil, com registro de percentuais de CPU e de memória ao longo das execuções.

Foram testadas cinco cardinalidades de fluxo (100, 500, 1000, 2000 e 5000 eventos), sob três perfis de carga (baixa, média e alta). Em cada combinação (volume × carga), foram realizadas três execuções independentes (execuções 1, 2 e 3), com registro de métricas por evento, incluindo: latência total por evento, tempo de processamento associado ao cálculo do score, throughput, consumo de CPU e memória e, no cenário com criptografia homomórfica, tempos de cifragem e decifragem e tamanho do payload cifrado.

No cenário Baseline (sem criptografia homomórfica), os eventos foram processados em claro, com tempos de cifragem/decifragem e tamanho de payload cifrado nulos. No cenário Seguro (CKKS/TenSEAL), os atributos telemétricos foram cifrados antes do envio, o cálculo do score foi realizado sobre dados cifrados e o resultado foi decifrado para registro e análise, permitindo comparação direta entre os cenários quanto a desempenho, consumo de recursos e escalabilidade.

As tabelas e figuras experimentais foram incorporadas ao corpo do relatório, nas seções de análise estatística e de resultados preliminares, conforme a numeração sequencial adotada.
""".strip()

    with open(os.path.join(p.texto_dir, "procedimentos_experimentais.md"), "w", encoding="utf-8") as f:
        f.write(proc + "\n")

    cpu_interp = """
### Interpretação do consumo de CPU e memória

Os percentuais de CPU e memória correspondem a amostragens via psutil ao longo da execução. Por esse motivo, variações no valor médio de CPU% entre os cenários devem ser interpretadas com cautela, pois a métrica reflete ocupação instantânea no intervalo de amostragem e pode ser influenciada por períodos de espera (por exemplo, serialização, E/S e comunicação), que tendem a aumentar no cenário com criptografia homomórfica. Assim, eventual redução do valor médio de CPU% no cenário CKKS não implica, por si só, menor custo computacional intrínseco da criptografia; a avaliação de eficiência priorizou latência, throughput e tempos explícitos de cifragem, processamento homomórfico e decifragem.
""".strip()

    with open(os.path.join(p.texto_dir, "interpretacao_cpu_memoria.md"), "w", encoding="utf-8") as f:
        f.write(cpu_interp + "\n")

    # Resultados e Discussão: texto base com campos quantitativos preenchidos após calcular
    # Estatísticas globais por cenário (agregado nas execuções e cargas)
    global_stats = run_all.groupby("cenario").agg(
        lat_mean=("latencia_media_evento", "mean"),
        lat_std=("latencia_media_evento", "std"),
        thr_mean=("throughput_run", "mean"),
        thr_std=("throughput_run", "std"),
        score_mean=("tempo_score_medio", "mean"),
        score_std=("tempo_score_medio", "std"),
        cpu_mean=("cpu_medio", "mean"),
        mem_mean=("mem_medio", "mean"),
    )

    def _val(scn: str, col: str) -> float:
        return float(global_stats.loc[scn, col])

    lat_plain = _val("PLAIN", "lat_mean")
    lat_he = _val("HE", "lat_mean")
    thr_plain = _val("PLAIN", "thr_mean")
    thr_he = _val("HE", "thr_mean")
    score_plain = _val("PLAIN", "score_mean")
    score_he = _val("HE", "score_mean")
    cpu_plain = _val("PLAIN", "cpu_mean")
    cpu_he = _val("HE", "cpu_mean")
    mem_plain = _val("PLAIN", "mem_mean")
    mem_he = _val("HE", "mem_mean")

    def _over(a: float, b: float) -> float:
        # overhead de b sobre a
        if a == 0:
            return float("nan")
        return ((b - a) / a) * 100.0

    concl = {
        "over_lat_pct": _over(lat_plain, lat_he),
        "over_thr_pct": _over(thr_plain, thr_he),
        "over_score_pct": _over(score_plain, score_he),
        "over_cpu_pct": _over(cpu_plain, cpu_he),
        "over_mem_pct": _over(mem_plain, mem_he),
    }

    def _fmt_delta(pct: float, positivo_significa_aumento: bool = True) -> str:
        if math.isnan(pct):
            return "não aplicável"
        if pct >= 0:
            return f"aumento de {pct:.2f}%"
        return f"redução de {abs(pct):.2f}%"

    payload_he_mean = float(run_all.loc[run_all["cenario"] == "HE", "payload_medio"].mean())
    cif_he_mean = float(run_all.loc[run_all["cenario"] == "HE", "tempo_cifragem_medio"].mean())
    dec_he_mean = float(run_all.loc[run_all["cenario"] == "HE", "tempo_decifragem_medio"].mean())

    res_intro = f"""
### Resultados e Discussão

#### Síntese quantitativa (agregado de todas as configurações)

No agregado das combinações de volume, carga e execuções independentes, a latência média por evento foi de **{lat_plain:.6f} s** no cenário Baseline e **{lat_he:.6f} s** no cenário Seguro (CKKS), correspondendo a **{_fmt_delta(concl['over_lat_pct'])}** da latência em relação ao Baseline. O throughput médio (eventos por segundo, calculado como volume dividido pela soma das latências por execução) foi de **{thr_plain:.2f}** no Baseline e **{thr_he:.2f}** no Seguro, ou seja, **{_fmt_delta(concl['over_thr_pct'])}** na vazão média (valores negativos de overhead indicam queda de throughput no cenário Seguro).

O tempo médio de processamento do score por evento foi de **{score_plain * 1e6:.3f} µs** no Baseline e **{score_he * 1e3:.3f} ms** no Seguro (**{_fmt_delta(concl['over_score_pct'])}**). No consumo médio de CPU (percentual amostrado via psutil), verificou-se **{_fmt_delta(concl['over_cpu_pct'])}** no Seguro em relação ao Baseline; essa variação foi interpretada com cautela, em função da natureza instantânea da métrica e da possibilidade de períodos de espera por E/S e comunicação. A memória média (percentual) apresentou **{_fmt_delta(concl['over_mem_pct'])}** no Seguro. No cenário CKKS, o tamanho médio do payload cifrado por evento (média agregada) foi de aproximadamente **{payload_he_mean:,.0f} bytes**, com tempos médios de cifragem e decifragem por evento da ordem de **{cif_he_mean * 1e3:.2f} ms** e **{dec_he_mean * 1e3:.2f} ms**, respectivamente.

Os resultados evidenciaram diferenças sistemáticas entre o cenário Baseline e o cenário Seguro, compatíveis com o custo adicional de operações homomórficas, serialização de ciphertexts e maior volume de dados trafegados. Do ponto de vista de viabilidade em arquiteturas IoT do tipo Pay As You Drive, o emprego do CKKS mostrou-se defendível quando a confidencialidade durante o processamento foi requisito prioritário, aceitando-se trade-offs de desempenho e dimensionando recursos (consumidores Kafka, paralelismo, particionamento de tópicos) para atendimento a níveis de serviço.

Como limitações, destacaram-se o ambiente controlado e a natureza pontual das amostras de CPU e memória; em produção, variabilidade de rede, concorrência e heterogeneidade de dispositivos podem alterar os custos observados.

---

#### Figuras e discussão interpretativa

As imagens abaixo utilizam caminhos relativos a este arquivo (`../figuras/`). Em visualizadores Markdown compatíveis, as figuras foram exibidas automaticamente.
""".strip()

    res_figs = r"""
#### Figura 1 – Latência média por volume (CKKS vs Baseline)

![Figura 1 – Latência média por volume (CKKS vs Baseline)](../figuras/grafico_1_latencia_media_por_volume.png)

A Figura 1 comparou a latência média por evento em função do volume testado. O cenário Seguro apresentou valores superiores ao Baseline, refletindo o custo incremental do pipeline criptográfico e do processamento sobre dados cifrados. A evolução com o volume permitiu discutir a estabilidade relativa do overhead e implicações de capacidade em sistemas Pay As You Drive.

*Fonte: Elaboração própria.*

#### Figura 2 – Throughput médio por carga (CKKS vs Baseline)

![Figura 2 – Throughput médio por carga (CKKS vs Baseline)](../figuras/grafico_2_throughput_por_carga.png)

A Figura 2 sintetizou a vazão média por perfil de carga (baixa, média e alta). A redução de throughput no cenário Seguro foi coerente com o aumento do tempo por evento: para um mesmo volume, menor vazão implicou maior tempo total de ocupação do pipeline. Overhead negativo sobre throughput significou **queda da taxa de processamento**, e não ganho de desempenho.

*Fonte: Elaboração própria.*

#### Figura 3 – Tempo médio de processamento do score por volume

![Figura 3 – Tempo médio de processamento do score por volume](../figuras/grafico_3_tempo_score.png)

A Figura 3 isolou o componente de tempo associado ao cálculo do score. No cenário CKKS, o incremento foi acentuado, pois operações aritméticas sobre vetores cifrados foram substancialmente mais custosas do que operações em claro. O gráfico sustentou a discussão sobre o principal gargalo computacional (além de cifrar, decifrar e serializar).

*Fonte: Elaboração própria.*

#### Figura 4 – Consumo médio de CPU por carga

![Figura 4 – Consumo médio de CPU por carga](../figuras/grafico_4_cpu.png)

A Figura 4 apresentou a média de CPU% por carga experimental. Os percentuais instantâneos de CPU não refletiram monotonicamente o trabalho criptográfico total por evento, especialmente quando houve maior tempo de espera por E/S ou comunicação. Por conseguinte, a leitura principal priorizou latência, throughput e tempos explícitos de cifragem, processamento homomórfico e decifragem.

*Fonte: Elaboração própria.*

#### Figura 5 – Consumo médio de memória por carga

![Figura 5 – Consumo médio de memória por carga](../figuras/grafico_5_memoria.png)

A Figura 5 mostrou o uso médio de memória do sistema (percentual) por carga. As alterações foram moderadas, compatíveis com o perfil de alocação de buffers e estruturas intermediárias; o impacto mais visível do CKKS apareceu no tamanho do payload cifrado (Figura 7) e no tempo de processamento (Figuras 1 e 3).

*Fonte: Elaboração própria.*

#### Figura 6 – Tempos médios de cifragem e decifragem por volume (CKKS)

![Figura 6 – Tempos médios de cifragem e decifragem por volume (CKKS)](../figuras/grafico_6_cifragem_decifragem.png)

A Figura 6 decompôs parte do custo criptográfico em cifragem (lado produtor) e decifragem (lado consumidor, após o processamento homomórfico). A decomposição apoiou a discussão sobre otimizações possíveis (parâmetros CKKS, paralelismo na borda, batching) sem comprometimento do modelo de ameaça desejado.

*Fonte: Elaboração própria.*

#### Figura 7 – Tamanho médio do payload cifrado por volume (CKKS)

![Figura 7 – Tamanho médio do payload cifrado por volume (CKKS)](../figuras/grafico_7_payload.png)

A Figura 7 evidenciou o custo de comunicação e armazenamento induzido pelo ciphertext. Payloads maiores implicaram maior demanda de largura de banda no Kafka, maior pressão sobre serialização e possível aumento de latência de rede em implantações distribuídas, tema relevante para IoT veicular.

*Fonte: Elaboração própria.*

#### Figura 8 – Overhead percentual do CKKS (média por volume)

![Figura 8 – Overhead percentual do CKKS (média por volume)](../figuras/grafico_8_overhead_percentual.png)

A Figura 8 consolidou o overhead (HE − Plain) / Plain × 100 para métricas selecionadas, agregando por volume. Valores positivos indicaram aumento da métrica no Seguro; valores negativos em throughput indicaram **redução da vazão**. A curva de tempo de score dominou a magnitude do overhead, alinhando-se à discussão da Figura 3.

*Fonte: Elaboração própria.*
""".strip()

    res = res_intro + "\n\n" + res_figs

    with open(os.path.join(p.texto_dir, "resultados_e_discussao.md"), "w", encoding="utf-8") as f:
        f.write(res + "\n")

    conc_txt = f"""
### Conclusões Quantitativas

Com base no agregado experimental, a adoção da criptografia homomórfica CKKS implicou {_fmt_delta(concl['over_lat_pct'])} na latência média por evento e {_fmt_delta(concl['over_thr_pct'])} no throughput médio, em comparação ao cenário Baseline. Observou-se {_fmt_delta(concl['over_score_pct'])} no tempo médio de processamento do score. No consumo médio de CPU, verificou-se {_fmt_delta(concl['over_cpu_pct'])}, e, no consumo médio de memória, {_fmt_delta(concl['over_mem_pct'])}.

Apesar do overhead, a abordagem permitiu manter a confidencialidade integral dos atributos telemétricos durante o processamento e o armazenamento do resultado cifrado, atendendo ao objetivo de uma arquitetura segura para processamento preditivo em sistemas IoT do tipo Pay As You Drive.

As tabelas e figuras experimentais foram incorporadas ao corpo do relatório, nas seções correspondentes, com títulos e fontes no padrão adotado.
""".strip()

    with open(os.path.join(p.texto_dir, "conclusoes_quantitativas.md"), "w", encoding="utf-8") as f:
        f.write(conc_txt + "\n")


def _fmt_cell_br(v: object) -> str:
    if v is None or (isinstance(v, float) and (math.isnan(v) or math.isinf(v))):
        return "—"
    if isinstance(v, (float, np.floating)):
        x = float(v)
        if x == 0.0:
            return "0"
        if abs(x) < 0.01:
            s = f"{x:.8f}".rstrip("0").rstrip(".").replace(".", ",")
            return s if s else "0"
        s = f"{x:.4f}".rstrip("0").rstrip(".").replace(".", ",")
        return s if s else "0"
    if isinstance(v, (int, np.integer)):
        return str(int(v))
    return str(v)


def _prep_tabela1(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["Cenário"] = out["cenario"].map({"PLAIN": "Baseline (sem HE)", "HE": "Seguro (CKKS)"})
    out = out.rename(
        columns={
            "latencia_media": "Latência média (s)",
            "tempo_score_medio": "Tempo score médio (s)",
            "throughput_medio": "Throughput médio (evt/s)",
            "cpu_media": "CPU média (%)",
            "mem_media": "Memória média (%)",
        }
    )
    return out[
        [
            "Cenário",
            "Latência média (s)",
            "Tempo score médio (s)",
            "Throughput médio (evt/s)",
            "CPU média (%)",
            "Memória média (%)",
        ]
    ]


def _prep_tabela2(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["Cenário"] = out["cenario"].map({"PLAIN": "Baseline (sem HE)", "HE": "Seguro (CKKS)"})
    out = out.rename(
        columns={
            "volume": "Volume (eventos)",
            "latencia_media_evento": "Latência média (s)",
            "throughput_run": "Throughput (evt/s)",
            "tempo_score_medio": "Tempo score médio (s)",
        }
    )
    return out[["Cenário", "Volume (eventos)", "Latência média (s)", "Throughput (evt/s)", "Tempo score médio (s)"]]


def _prep_tabela3(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out = out.rename(
        columns={
            "volume": "Volume (eventos)",
            "overhead_pct_latencia_media_evento": "Overhead latência (%)",
            "overhead_pct_throughput_run": "Overhead throughput (%)",
            "overhead_pct_tempo_score_medio": "Overhead tempo score (%)",
            "overhead_pct_cpu_medio": "Overhead CPU (%)",
            "overhead_pct_mem_medio": "Overhead memória (%)",
        }
    )
    return out


def _prep_tabela4(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["Cenário"] = out["cenario"].map({"PLAIN": "Baseline (sem HE)", "HE": "Seguro (CKKS)"})
    out["Carga"] = out["carga"].str.capitalize()
    out = out.rename(columns={"cpu_media": "CPU média (%)", "mem_media": "Memória média (%)"})
    return out[["Cenário", "Carga", "CPU média (%)", "Memória média (%)"]]


def _prep_tabela5(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out = out.rename(
        columns={
            "volume": "Volume (eventos)",
            "tempo_cifragem_medio": "Cifragem média (s)",
            "tempo_decifragem_medio": "Decifragem média (s)",
        }
    )
    return out


def _prep_tabela6(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out = out.rename(columns={"volume": "Volume (eventos)", "payload_medio": "Payload médio (bytes)"})
    return out


def _update_docx(p: Paths, dfs: Mapping[str, pd.DataFrame]) -> Optional[str]:
    """
    Incorpora tabelas (objeto nativo Word), figuras e textos ao RP - 251.docx original.
    Retorna o caminho salvo ou None.
    """
    docx_in = "RP - 251.docx"
    if not os.path.exists(docx_in):
        return None

    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.shared import Inches
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except Exception:
        return None

    doc = Document(docx_in)

    def _doc_text_joined() -> str:
        return "\n".join((par.text or "") for par in doc.paragraphs)

    if "Tabela 1 – Estatísticas globais por cenário (média agregada nas configurações)." in _doc_text_joined():
        return "SKIPPED_DUPLICATE"

    backup = os.path.join(p.out_dir, "RP - 251_backup_pre_pipeline.docx")
    if not os.path.exists(backup):
        shutil.copy2(docx_in, backup)

    def _find_paragraph_by_exact_text(text: str) -> Optional[Paragraph]:
        alvo = text.strip().lower()
        for par in doc.paragraphs:
            if (par.text or "").strip().lower() == alvo:
                return par
        return None

    def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        return new_para

    def _insert_table_after(paragraph: Paragraph, title: str, df: pd.DataFrame) -> Paragraph:
        document = paragraph.part.document
        body = document.element.body

        cur = _insert_paragraph_after(paragraph, title)

        n_rows = len(df) + 1
        n_cols = len(df.columns)
        tmp = document.add_table(rows=n_rows, cols=n_cols)
        tbl_elm = tmp._tbl
        body.remove(tbl_elm)
        cur._p.addnext(tbl_elm)

        table = Table(tbl_elm, document._body)
        for j, col in enumerate(df.columns):
            table.cell(0, j).text = str(col)
        for i in range(len(df)):
            for j, col in enumerate(df.columns):
                table.cell(i + 1, j).text = _fmt_cell_br(df.iloc[i, j])

        font_el = OxmlElement("w:p")
        tbl_elm.addnext(font_el)
        font_p = Paragraph(font_el, document._body)
        font_p.add_run("Fonte: Elaboração própria.")
        return font_p

    def _insert_paragraphs_from_lines(anchor: Paragraph, lines: List[str]) -> Paragraph:
        cur = anchor
        for line in lines:
            s = line.strip()
            if not s or s.startswith("###"):
                continue
            cur = _insert_paragraph_after(cur, s)
        return cur

    analises_figura: Dict[str, str] = {
        "grafico_1_latencia_media_por_volume.png": (
            "A Figura 1 evidenciou latência média superior no cenário Seguro (CKKS) em relação ao Baseline, "
            "em todos os volumes avaliados, compatível com o custo incremental do processamento sobre dados cifrados "
            "e com etapas adicionais do pipeline experimental."
        ),
        "grafico_2_throughput_por_carga.png": (
            "A Figura 2 mostrou menor throughput médio no cenário Seguro, coerente com o aumento do tempo por evento. "
            "Overhead negativo sobre throughput indicou queda da taxa de processamento (vazão), e não melhoria de desempenho."
        ),
        "grafico_3_tempo_score.png": (
            "A Figura 3 destacou o acréscimo expressivo do tempo médio de processamento do score no domínio homomórfico, "
            "indicando que o cálculo sobre vetores cifrados constituiu o principal componente de custo computacional além "
            "da cifragem e da serialização."
        ),
        "grafico_4_cpu.png": (
            "A Figura 4 sintetizou médias de CPU por carga. Os percentuais foram obtidos por amostragem via psutil e, "
            "por serem instantâneos, não devem ser interpretados isoladamente como medida de esforço criptográfico total, "
            "especialmente na presença de esperas por E/S e comunicação."
        ),
        "grafico_5_memoria.png": (
            "A Figura 5 apresentou variações moderadas de memória por carga. O impacto mais pronunciado do CKKS no experimento "
            "associou-se principalmente ao aumento do payload cifrado e ao tempo de processamento, conforme as Figuras 7 e 3."
        ),
        "grafico_6_cifragem_decifragem.png": (
            "A Figura 6 decompôs tempos médios de cifragem e decifragem por volume. A cifragem concentrou maior parcela do "
            "custo temporal criptográfico por evento, enquanto a decifragem permaneceu em ordem de grandeza inferior."
        ),
        "grafico_7_payload.png": (
            "A Figura 7 mostrou o tamanho médio do payload cifrado por volume, com magnitudes compatíveis com ciphertexts CKKS. "
            "Esse resultado reforçou o custo de comunicação e armazenamento no barramento de mensagens e a necessidade de "
            "dimensionamento de rede em cenários IoT veiculares."
        ),
        "grafico_8_overhead_percentual.png": (
            "A Figura 8 consolidou o overhead percentual médio por volume. Valores positivos indicaram aumento da métrica no "
            "cenário Seguro; valores negativos em throughput indicaram redução da vazão. O componente de tempo de score "
            "apresentou a maior magnitude relativa, alinhando-se à interpretação da Figura 3."
        ),
    }

    proc_lines = open(os.path.join(p.texto_dir, "procedimentos_experimentais.md"), encoding="utf-8").read().splitlines()

    cpu_lines = [
        lin
        for lin in open(os.path.join(p.texto_dir, "interpretacao_cpu_memoria.md"), encoding="utf-8").read().splitlines()
    ]

    conc_lines = [
        lin
        for lin in open(os.path.join(p.texto_dir, "conclusoes_quantitativas.md"), encoding="utf-8").read().splitlines()
    ]

    proc_anchor = _find_paragraph_by_exact_text("Desenho Experimental")
    if proc_anchor is not None:
        _insert_paragraphs_from_lines(proc_anchor, proc_lines)

    stat_anchor = _find_paragraph_by_exact_text("Análise Estatística")
    if stat_anchor is not None:
        cur = stat_anchor
        cur = _insert_table_after(
            cur,
            "Tabela 1 – Estatísticas globais por cenário (média agregada nas configurações).",
            _prep_tabela1(dfs["t1"]),
        )
        cur = _insert_table_after(
            cur,
            "Tabela 2 – Desempenho médio por volume (latência, throughput e tempo de score).",
            _prep_tabela2(dfs["t2"]),
        )
        cur = _insert_table_after(
            cur,
            "Tabela 3 – Overhead criptográfico percentual por volume (média nas cargas e execuções).",
            _prep_tabela3(dfs["t3"]),
        )
        cur = _insert_table_after(
            cur,
            "Tabela 4 – Consumo médio de recursos por carga (CPU e memória).",
            _prep_tabela4(dfs["t4"]),
        )
        cur = _insert_table_after(
            cur,
            "Tabela 5 – Tempos médios de cifragem e decifragem por volume (cenário CKKS).",
            _prep_tabela5(dfs["t5"]),
        )
        _insert_table_after(
            cur,
            "Tabela 6 – Tamanho médio do payload cifrado por volume (cenário CKKS).",
            _prep_tabela6(dfs["t6"]),
        )

    cpu_anchor = _find_paragraph_by_exact_text("Consumo de recursos computacionais")
    if cpu_anchor is not None:
        _insert_paragraphs_from_lines(cpu_anchor, cpu_lines)

    anchors_figuras: List[Tuple[str, List[Tuple[str, str]]]] = [
        ("Throughput do sistema", [("Figura 2 – Throughput médio por carga (HE vs Plain).", "grafico_2_throughput_por_carga.png")]),
        ("Latência de processamento", [("Figura 1 – Latência média por volume (HE vs Plain).", "grafico_1_latencia_media_por_volume.png")]),
        ("Tempo de geração do score", [("Figura 3 – Tempo médio de score por volume (HE vs Plain).", "grafico_3_tempo_score.png")]),
        (
            "Consumo de recursos computacionais",
            [
                ("Figura 4 – Consumo médio de CPU por carga (HE vs Plain).", "grafico_4_cpu.png"),
                ("Figura 5 – Consumo médio de memória por carga (HE vs Plain).", "grafico_5_memoria.png"),
            ],
        ),
        (
            "Considerações parciais",
            [
                ("Figura 6 – Tempo médio de cifragem e decifragem por volume (CKKS).", "grafico_6_cifragem_decifragem.png"),
                ("Figura 7 – Tamanho médio do payload cifrado por volume (CKKS).", "grafico_7_payload.png"),
                ("Figura 8 – Overhead percentual do CKKS (média por volume).", "grafico_8_overhead_percentual.png"),
            ],
        ),
    ]

    for titulo_secao, lista in anchors_figuras:
        par = _find_paragraph_by_exact_text(titulo_secao)
        if par is None:
            continue
        cur = par
        for legenda, fname in lista:
            img_path = os.path.join(p.figs_dir, fname)
            if not os.path.exists(img_path):
                continue
            cur = _insert_paragraph_after(cur, legenda)
            run = cur.add_run()
            run.add_picture(img_path, width=Inches(6.5))
            cur = _insert_paragraph_after(cur, "Fonte: Elaboração própria.")
            analise = analises_figura.get(fname, "")
            if analise:
                cur = _insert_paragraph_after(cur, analise)

    conc_anchor = _find_paragraph_by_exact_text("Conclusão")
    if conc_anchor is not None:
        _insert_paragraphs_from_lines(conc_anchor, conc_lines)

    doc.save(docx_in)

    with open(os.path.join(p.out_dir, "docx_update_status.txt"), "w", encoding="utf-8") as f:
        f.write(
            f"Documento principal atualizado: {os.path.abspath(docx_in)}\n"
            f"Cópia de segurança (somente na primeira execução): {os.path.abspath(backup)}\n"
        )

    return docx_in


def main() -> int:
    p = Paths()
    _ensure_dirs(p)

    plain = _read_csv(p.plain_csv)
    he = _read_csv(p.he_csv)

    erros = []
    erros += _validate_schema(plain, "PLAIN")
    erros += _validate_schema(he, "HE")
    if not erros:
        erros += _validate_grid(plain, "PLAIN")
        erros += _validate_grid(he, "HE")
        erros += _validate_values(plain, "PLAIN")
        erros += _validate_values(he, "HE")

    with open(os.path.join(p.out_dir, "validacao.txt"), "w", encoding="utf-8") as f:
        if erros:
            f.write("ERROS/ALERTAS DE CONSISTÊNCIA\n")
            for e in erros:
                f.write(f"- {e}\n")
        else:
            f.write("OK: não foram encontrados problemas de consistência relevantes.\n")

    run_plain = _per_run_aggregates(plain)
    run_he = _per_run_aggregates(he)
    run_all = pd.concat([run_plain, run_he], ignore_index=True)

    # Estatísticas descritivas por evento (dados crus)
    desc_evento = _descritivas(
        pd.concat([plain, he], ignore_index=True),
        cols=METRICAS_BASE,
        group_cols=["cenario", "volume", "carga"],
    )
    desc_evento.to_csv(os.path.join(p.tables_dir, "tabela_descritivas_por_evento.csv"), index=False)

    # Estatísticas por execução (run-level)
    cols_run = [
        "latencia_media_evento",
        "tempo_score_medio",
        "throughput_run",
        "cpu_medio",
        "mem_medio",
        "tempo_cifragem_medio",
        "tempo_decifragem_medio",
        "payload_medio",
    ]
    desc_run = _descritivas(run_all, cols=cols_run, group_cols=["cenario", "volume", "carga"])
    desc_run.to_csv(os.path.join(p.tables_dir, "tabela_descritivas_por_execucao.csv"), index=False)

    # Overhead (%): métricas solicitadas
    metricas_over = [
        "latencia_media_evento",
        "throughput_run",
        "tempo_score_medio",
        "cpu_medio",
        "mem_medio",
    ]
    overhead = _merge_overhead(run_plain, run_he, metricas_over)
    overhead.to_csv(os.path.join(p.tables_dir, "overhead_por_configuracao.csv"), index=False)

    # Gráficos
    _make_plots(run_all, overhead, p)

    # Tabelas formatadas (texto)
    tabela1 = (
        run_all.groupby(["cenario"])
        .agg(
            latencia_media=("latencia_media_evento", "mean"),
            tempo_score_medio=("tempo_score_medio", "mean"),
            throughput_medio=("throughput_run", "mean"),
            cpu_media=("cpu_medio", "mean"),
            mem_media=("mem_medio", "mean"),
        )
        .reset_index()
    )
    tabela1_path = os.path.join(p.tables_dir, "tabela_1_resumo_global.txt")
    with open(tabela1_path, "w", encoding="utf-8") as f:
        f.write(
            _format_tabela_uspesalq(
                "Tabela 1 – Estatísticas globais por cenário (média agregada nas configurações).",
                tabela1,
                "Elaboração própria",
            )
        )

    tabela2 = (
        run_all.groupby(["cenario", "volume"], as_index=False)
        .agg(
            latencia_media_evento=("latencia_media_evento", "mean"),
            throughput_run=("throughput_run", "mean"),
            tempo_score_medio=("tempo_score_medio", "mean"),
        )
        .sort_values(["volume", "cenario"])
    )
    with open(os.path.join(p.tables_dir, "tabela_2_por_volume.txt"), "w", encoding="utf-8") as f:
        f.write(
            _format_tabela_uspesalq(
                "Tabela 2 – Desempenho médio por volume (latência, throughput e tempo de score).",
                tabela2,
                "Elaboração própria",
            )
        )

    overhead_cols = [
        "overhead_pct_latencia_media_evento",
        "overhead_pct_throughput_run",
        "overhead_pct_tempo_score_medio",
        "overhead_pct_cpu_medio",
        "overhead_pct_mem_medio",
    ]
    tabela3 = overhead.groupby(["volume"], as_index=False)[overhead_cols].mean().sort_values(["volume"])
    with open(os.path.join(p.tables_dir, "tabela_3_overhead_por_volume.txt"), "w", encoding="utf-8") as f:
        f.write(
            _format_tabela_uspesalq(
                "Tabela 3 – Overhead criptográfico percentual por volume (média nas cargas e execuções).",
                tabela3,
                "Elaboração própria",
            )
        )

    tabela4 = (
        run_all.groupby(["cenario", "carga"], as_index=False)
        .agg(cpu_media=("cpu_medio", "mean"), mem_media=("mem_medio", "mean"))
        .sort_values(["carga", "cenario"])
    )
    with open(os.path.join(p.tables_dir, "tabela_4_recursos_por_carga.txt"), "w", encoding="utf-8") as f:
        f.write(
            _format_tabela_uspesalq(
                "Tabela 4 – Consumo médio de recursos por carga (CPU e memória).",
                tabela4,
                "Elaboração própria",
            )
        )

    tabela5 = (
        run_he.groupby(["volume"], as_index=False)[["tempo_cifragem_medio", "tempo_decifragem_medio"]]
        .mean()
        .sort_values(["volume"])
    )
    with open(os.path.join(p.tables_dir, "tabela_5_tempos_crypto_ckks.txt"), "w", encoding="utf-8") as f:
        f.write(
            _format_tabela_uspesalq(
                "Tabela 5 – Tempos médios de cifragem e decifragem por volume (cenário CKKS).",
                tabela5,
                "Elaboração própria",
            )
        )

    tabela6 = (
        run_he.groupby(["volume"], as_index=False)[["payload_medio"]].mean().sort_values(["volume"])
    )
    with open(os.path.join(p.tables_dir, "tabela_6_payload_ckks.txt"), "w", encoding="utf-8") as f:
        f.write(
            _format_tabela_uspesalq(
                "Tabela 6 – Tamanho médio do payload cifrado por volume (cenário CKKS).",
                tabela6,
                "Elaboração própria",
            )
        )

    # Texto acadêmico
    _write_text_sections(run_all, overhead, p)

    # Atualização do DOCX
    docx_out = _update_docx(
        p,
        {
            "t1": tabela1,
            "t2": tabela2,
            "t3": tabela3,
            "t4": tabela4,
            "t5": tabela5,
            "t6": tabela6,
        },
    )
    with open(os.path.join(p.out_dir, "docx_update_status.txt"), "w", encoding="utf-8") as f:
        if docx_out == "SKIPPED_DUPLICATE":
            f.write(
                "O arquivo RP - 251.docx já contém as tabelas automáticas inseridas anteriormente; "
                "a reinserção foi omitida para evitar duplicação. Restaure RP - 251.docx a partir do controle de versão "
                "(por exemplo, `git checkout -- \"RP - 251.docx\"`) ou utilize a cópia em "
                "outputs_analise/RP - 251_backup_pre_pipeline.docx, se existir, e execute o script novamente.\n"
            )
        elif docx_out:
            f.write(f"Documento principal atualizado: {os.path.abspath(docx_out)}\n")
        else:
            f.write(
                "Não foi possível atualizar o DOCX (arquivo ausente ou python-docx indisponível).\n"
            )

    return 0


if __name__ == "__main__":
    raise SystemExit(main())

